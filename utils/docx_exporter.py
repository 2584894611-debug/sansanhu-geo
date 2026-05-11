# -*- coding: utf-8 -*-
"""
好易易GEO - Word文档导出模块
将Markdown格式的报告转换为Word文档
"""

import re
from typing import Optional, List, Dict
from pathlib import Path


class MarkdownToDocxConverter:
    """
    Markdown转Word文档转换器
    支持标题、段落、列表、表格等常见格式
    """
    
    def __init__(self):
        """初始化转换器"""
        self.current_style = None
        
    def convert(self, markdown_text: str) -> 'Document':
        """
        将Markdown文本转换为python-docx文档对象
        
        Args:
            markdown_text: Markdown格式文本
        
        Returns:
            python-docx Document对象
        """
        try:
            from docx import Document
            from docx.shared import Pt, RGBColor, Inches
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            from docx.enum.table import WD_TABLE_ALIGNMENT
        except ImportError:
            raise ImportError(
                "请先安装 python-docx 库: pip install python-docx"
            )
        
        doc = Document()
        
        # 设置默认字体
        style = doc.styles['Normal']
        font = style.font
        font.name = '微软雅黑'
        font.size = Pt(11)
        
        # 处理标题样式
        try:
            heading_style = doc.styles['Heading 1']
            heading_style.font.name = '微软雅黑'
            heading_style.font.size = Pt(18)
            heading_style.font.bold = True
        except:
            pass
        
        # 解析Markdown并添加到文档
        lines = markdown_text.split('\n')
        self._parse_markdown(doc, lines)
        
        return doc
    
    def _parse_markdown(self, doc, lines: List[str]):
        """解析Markdown内容"""
        from docx import Document
        from docx.shared import Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        
        i = 0
        in_code_block = False
        code_content = []
        table_rows = []
        in_table = False
        in_list = False
        
        while i < len(lines):
            line = lines[i]
            
            # 代码块
            if line.strip().startswith('```'):
                if in_code_block:
                    # 结束代码块
                    self._add_code_block(doc, '\n'.join(code_content))
                    code_content = []
                    in_code_block = False
                else:
                    # 开始代码块
                    in_code_block = True
                i += 1
                continue
            
            if in_code_block:
                code_content.append(line)
                i += 1
                continue
            
            # 表格处理
            if '|' in line and line.strip().startswith('|'):
                if not in_table:
                    in_table = True
                    table_rows = []
                # 解析表格行
                cells = [cell.strip() for cell in line.split('|')]
                cells = [c for c in cells if c]  # 移除空单元格
                if cells and not all(c.strip().startswith('-') for c in cells if c):
                    table_rows.append(cells)
                i += 1
                continue
            else:
                if in_table and table_rows:
                    self._add_table(doc, table_rows)
                    table_rows = []
                in_table = False
            
            # 标题
            if line.startswith('# '):
                text = line[2:].strip()
                self._add_heading(doc, text, 1)
            elif line.startswith('## '):
                text = line[3:].strip()
                self._add_heading(doc, text, 2)
            elif line.startswith('### '):
                text = line[4:].strip()
                self._add_heading(doc, text, 3)
            elif line.startswith('#### '):
                text = line[5:].strip()
                self._add_heading(doc, text, 4)
            
            # 分隔线
            elif line.strip() in ['---', '***', '___']:
                self._add_horizontal_rule(doc)
            
            # 引用块
            elif line.startswith('> '):
                text = line[2:].strip()
                self._add_blockquote(doc, text)
            
            # 列表项
            elif line.strip().startswith(('- ', '* ', '+ ')):
                text = line.strip()[2:].strip()
                self._add_list_item(doc, text, 'unordered')
            elif line.strip()[0:2].isdigit() and line.strip()[1] == '.':
                # 有序列表
                text = re.sub(r'^\d+\.\s*', '', line.strip())
                self._add_list_item(doc, text, 'ordered')
            
            # 段落
            elif line.strip():
                # 处理行内格式
                formatted_text = self._process_inline_formatting(line.strip())
                self._add_paragraph(doc, formatted_text)
            
            i += 1
        
        # 处理未结束的表格
        if in_table and table_rows:
            self._add_table(doc, table_rows)
    
    def _add_heading(self, doc, text: str, level: int):
        """添加标题"""
        from docx.shared import Pt
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        
        heading = doc.add_heading(text, level=level)
        heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
        
        # 设置字体
        for run in heading.runs:
            run.font.name = '微软雅黑'
            if level <= 2:
                run.font.bold = True
    
    def _add_paragraph(self, doc, text: str):
        """添加段落"""
        from docx.shared import Pt
        
        para = doc.add_paragraph()
        self._apply_inline_formatting(para, text)
        return para
    
    def _add_code_block(self, doc, code: str):
        """添加代码块"""
        from docx.shared import Pt, RGBColor
        from docx.oxml.ns import qn
        
        para = doc.add_paragraph()
        run = para.add_run(code)
        run.font.name = 'Consolas'
        run.font.size = Pt(9)
        # 设置浅灰色背景（通过XML设置）
        
        # 添加边框
        para.paragraph_format.left_indent = Pt(20)
        para.paragraph_format.space_before = Pt(6)
        para.paragraph_format.space_after = Pt(6)
    
    def _add_blockquote(self, doc, text: str):
        """添加引用块"""
        from docx.shared import Pt
        from docx.oxml.ns import qn
        
        para = doc.add_paragraph()
        para.paragraph_format.left_indent = Pt(20)
        
        run = para.add_run(text)
        run.font.name = '微软雅黑'
        run.font.size = Pt(10)
        run.font.italic = True
        run.font.color.rgb = None  # 保持默认颜色
    
    def _add_list_item(self, doc, text: str, list_type: str = 'unordered'):
        """添加列表项"""
        from docx.shared import Pt
        
        para = doc.add_paragraph(style='List Bullet' if list_type == 'unordered' else 'List Number')
        self._apply_inline_formatting(para, text)
    
    def _add_table(self, doc, rows: List[List[str]]):
        """添加表格"""
        from docx.shared import Pt, Cm
        from docx.enum.table import WD_TABLE_ALIGNMENT
        
        if not rows:
            return
        
        # 创建表格
        table = doc.add_table(rows=len(rows), cols=len(rows[0]))
        table.style = 'Table Grid'
        
        for i, row_data in enumerate(rows):
            row = table.rows[i]
            for j, cell_text in enumerate(row_data):
                cell = row.cells[j]
                cell.text = cell_text
                
                # 表头样式
                if i == 0:
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.font.bold = True
                            run.font.name = '微软雅黑'
                else:
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.font.name = '微软雅黑'
        
        doc.add_paragraph()  # 表格后添加空行
    
    def _add_horizontal_rule(self, doc):
        """添加分隔线"""
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement
        
        para = doc.add_paragraph()
        p = para._p
        pPr = p.get_or_add_pPr()
        
        # 添加底边框
        pBdr = OxmlElement('w:pBdr')
        bottom = OxmlElement('w:bottom')
        bottom.set(qn('w:val'), 'single')
        bottom.set(qn('w:sz'), '6')
        bottom.set(qn('w:space'), '1')
        bottom.set(qn('w:color'), 'auto')
        pBdr.append(bottom)
        pPr.append(pBdr)
    
    def _process_inline_formatting(self, text: str) -> str:
        """
        处理行内格式标记
        返回处理后的文本（保留星号等标记供后续处理）
        """
        # 简化处理：移除行内代码标记但保留内容
        # 实际应用会在 _apply_inline_formatting 中处理
        return text
    
    def _apply_inline_formatting(self, para, text: str):
        """
        应用行内格式（粗体、斜体等）
        """
        from docx.shared import Pt, RGBColor
        
        # 处理加粗 **text**
        parts = re.split(r'(\*\*[^*]+\*\*|\*[^*]+\*|`[^`]+`)', text)
        
        for part in parts:
            run = para.add_run()
            run.font.name = '微软雅黑'
            run.font.size = Pt(11)
            
            if part.startswith('**') and part.endswith('**'):
                # 加粗
                run.text = part[2:-2]
                run.font.bold = True
            elif part.startswith('*') and part.endswith('*'):
                # 斜体
                run.text = part[1:-1]
                run.font.italic = True
            elif part.startswith('`') and part.endswith('`'):
                # 行内代码
                run.text = part[1:-1]
                run.font.name = 'Consolas'
                run.font.size = Pt(10)
            else:
                # 普通文本
                run.text = part


def markdown_to_docx(
    markdown_text: str,
    output_path: str,
    title: Optional[str] = None,
    author: Optional[str] = None
) -> bool:
    """
    将Markdown文本转换为Word文档
    
    Args:
        markdown_text: Markdown格式文本
        output_path: 输出文件路径
        title: 文档标题（可选）
        author: 作者（可选）
    
    Returns:
        转换是否成功
    """
    try:
        from docx import Document
        from docx.shared import Pt, Inches
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.enum.style import WD_STYLE_TYPE
        
        # 创建转换器
        converter = MarkdownToDocxConverter()
        doc = converter.convert(markdown_text)
        
        # 设置文档属性
        if title:
            doc.core_properties.title = title
        if author:
            doc.core_properties.author = author
        
        # 设置页面边距
        for section in doc.sections:
            section.top_margin = Inches(1)
            section.bottom_margin = Inches(1)
            section.left_margin = Inches(1.25)
            section.right_margin = Inches(1.25)
        
        # 保存文档
        doc.save(output_path)
        return True
        
    except ImportError as e:
        print(f"错误: {e}")
        print("请运行: pip install python-docx")
        return False
    except Exception as e:
        print(f"转换失败: {e}")
        return False


def markdown_file_to_docx(
    input_path: str,
    output_path: str,
    title: Optional[str] = None,
    author: Optional[str] = None
) -> bool:
    """
    将Markdown文件转换为Word文档
    
    Args:
        input_path: 输入Markdown文件路径
        output_path: 输出Word文件路径
        title: 文档标题（可选）
        author: 作者（可选）
    
    Returns:
        转换是否成功
    """
    try:
        with open(input_path, 'r', encoding='utf-8') as f:
            markdown_text = f.read()
        
        return markdown_to_docx(markdown_text, output_path, title, author)
        
    except FileNotFoundError:
        print(f"错误: 文件未找到 - {input_path}")
        return False
    except Exception as e:
        print(f"转换失败: {e}")
        return False


def add_cover_page(
    doc,
    title: str,
    subtitle: str = "",
    brand_name: str = "",
    date: str = "",
    author: str = ""
):
    """
    添加封面页到Word文档
    
    Args:
        doc: Document对象
        title: 主标题
        subtitle: 副标题
        brand_name: 品牌名称
        date: 日期
        author: 作者
    """
    from docx.shared import Pt, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    
    # 封面页
    doc.add_page_break()
    
    # 主标题
    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title_para.add_run(title)
    run.font.name = '微软雅黑'
    run.font.size = Pt(28)
    run.font.bold = True
    
    # 副标题
    if subtitle:
        sub_para = doc.add_paragraph()
        sub_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = sub_para.add_run(subtitle)
        run.font.name = '微软雅黑'
        run.font.size = Pt(14)
    
    # 空行
    doc.add_paragraph()
    doc.add_paragraph()
    
    # 品牌名称
    if brand_name:
        brand_para = doc.add_paragraph()
        brand_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = brand_para.add_run(f"品牌：{brand_name}")
        run.font.name = '微软雅黑'
        run.font.size = Pt(12)
    
    # 日期
    if date:
        date_para = doc.add_paragraph()
        date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = date_para.add_run(f"日期：{date}")
        run.font.name = '微软雅黑'
        run.font.size = Pt(11)
    
    # 作者
    if author:
        author_para = doc.add_paragraph()
        author_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = author_para.add_run(f"作者：{author}")
        run.font.name = '微软雅黑'
        run.font.size = Pt(11)
    
    # 添加页眉页脚
    doc.add_page_break()


def add_toc(doc):
    """
    添加目录占位符到文档
    
    Args:
        doc: Document对象
    """
    from docx.shared import Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    
    doc.add_paragraph()
    
    toc_title = doc.add_paragraph()
    toc_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = toc_title.add_run("目  录")
    run.font.name = '微软雅黑'
    run.font.size = Pt(16)
    run.font.bold = True
    
    doc.add_paragraph()
    
    toc_entries = [
        "一、AI搜索表现",
        "二、品牌健康度",
        "三、竞品对比",
        "四、优化建议",
        "五、内容策略"
    ]
    
    for entry in toc_entries:
        para = doc.add_paragraph()
        run = para.add_run(entry)
        run.font.name = '微软雅黑'
        run.font.size = Pt(12)
        para.paragraph_format.left_indent = Inches(2)
    
    doc.add_page_break()


def create_full_report_docx(
    markdown_text: str,
    output_path: str,
    report_title: str,
    brand_name: str,
    report_date: str = "",
    include_cover: bool = True,
    include_toc: bool = True
) -> bool:
    """
    创建完整的GEO分析报告Word文档
    
    Args:
        markdown_text: Markdown格式的报告内容
        output_path: 输出文件路径
        report_title: 报告标题
        brand_name: 品牌名称
        report_date: 报告日期
        include_cover: 是否包含封面
        include_toc: 是否包含目录
    
    Returns:
        创建是否成功
    """
    try:
        from docx import Document
        from docx.shared import Pt, Inches
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        
        # 创建转换器
        converter = MarkdownToDocxConverter()
        doc = converter.convert(markdown_text)
        
        # 设置文档属性
        doc.core_properties.title = report_title
        doc.core_properties.author = "好易易GEO"
        
        # 设置页面边距
        for section in doc.sections:
            section.top_margin = Inches(1)
            section.bottom_margin = Inches(1)
            section.left_margin = Inches(1.25)
            section.right_margin = Inches(1.25)
        
        # 添加页眉页脚
        for section in doc.sections:
            # 页眉
            header = section.header
            header_para = header.paragraphs[0]
            header_para.text = f"好易易GEO | {brand_name} GEO分析报告"
            header_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # 页脚
            footer = section.footer
            footer_para = footer.paragraphs[0]
            footer_para.text = "机密 | 仅供内部使用"
            footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # 保存文档
        doc.save(output_path)
        return True
        
    except ImportError as e:
        print(f"错误: {e}")
        print("请运行: pip install python-docx")
        return False
    except Exception as e:
        print(f"创建报告失败: {e}")
        return False


if __name__ == '__main__':
    # 测试转换功能
    
    # 示例Markdown内容
    test_markdown = """# 测试品牌 GEO分析报告

## 一、AI搜索表现

本报告分析了测试品牌在各AI搜索平台的表现情况。

### 1.1 平台提及详情

| 平台 | 提及次数 | 权重 |
|:---|:---:|:---:|
| 小红书 | 85 | ⭐⭐⭐ |
| 知乎 | 72 | ⭐⭐⭐ |
| 微信公众号 | 65 | ⭐⭐⭐ |

### 1.2 可见度分析

- 品牌可见度较高
- 建议加强内容优化

## 二、品牌健康度

### GEO评分

**GEO健康评分**：78/100

```
[████████████████░░░░░░░░░░░░░░░░░░░░░░░░░░] 78%
```

## 三、优化建议

1. 加强小红书内容布局
2. 增加知乎高质量回答
3. 优化微信公众号内容策略

> 这是引用内容
"""
    
    # 测试Markdown转Docx
    output_file = "test_report.docx"
    success = markdown_to_docx(test_markdown, output_file, title="测试报告")
    
    if success:
        print(f"✅ 报告已生成: {output_file}")
    else:
        print("❌ 报告生成失败")
