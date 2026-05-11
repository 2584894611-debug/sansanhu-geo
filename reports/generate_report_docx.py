#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
生成武汉国际广场GEO报告的精美Word文档
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def set_cell_shading(cell, color):
    """设置单元格背景色"""
    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:fill'), color)
    cell._tc.get_or_add_tcPr().append(shading_elm)

def add_horizontal_line(paragraph):
    """添加水平分隔线"""
    p = paragraph._p
    pPr = p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), '666666')
    pBdr.append(bottom)
    pPr.append(pBdr)

def add_run_with_color(paragraph, text, color=None, bold=False, size=None):
    """添加带样式的文本"""
    run = paragraph.add_run(text)
    if color:
        run.font.color.rgb = RGBColor(*color)
    if bold:
        run.font.bold = bold
    if size:
        run.font.size = Pt(size)
    return run

# 创建文档
doc = Document()

# 设置默认字体
style = doc.styles['Normal']
style.font.name = '微软雅黑'
style.font.size = Pt(11)
style._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

# ========== 标题 ==========
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
title_run = title.add_run('武汉国际广场 GEO分析报告')
title_run.font.size = Pt(22)
title_run.font.bold = True
title_run.font.color.rgb = RGBColor(51, 51, 51)
title_run.font.name = '微软雅黑'
title_run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
title.paragraph_format.space_after = Pt(12)

# 日期和行业信息
info = doc.add_paragraph()
info.alignment = WD_ALIGN_PARAGRAPH.CENTER
info_run = info.add_run('📅 报告日期：2026年05月09日\n🏢 所属行业：商业地产/本地生活')
info_run.font.size = Pt(11)
info_run.font.color.rgb = RGBColor(102, 102, 102)
info.paragraph_format.space_after = Pt(12)

# 分隔线
line_para = doc.add_paragraph()
add_horizontal_line(line_para)
line_para.paragraph_format.space_after = Pt(12)

# ========== 执行摘要 ==========
section_title = doc.add_paragraph()
st_run = section_title.add_run('📊 执行摘要')
st_run.font.size = Pt(16)
st_run.font.bold = True
st_run.font.color.rgb = RGBColor(51, 51, 51)
section_title.paragraph_format.space_before = Pt(6)
section_title.paragraph_format.space_after = Pt(8)

# 执行摘要内容
summary_intro = doc.add_paragraph()
summary_intro.add_run('通过对武汉国际广场在主流AI搜索平台的表现进行综合分析，得出以下核心发现：')
summary_intro.paragraph_format.space_after = Pt(12)

# 评分表格
table = doc.add_table(rows=1, cols=1)
table.style = 'Table Grid'
cell = table.rows[0].cells[0]
cell.text = ''
cell_p = cell.paragraphs[0]

# GEO健康评分 - 醒目样式
geo_score = cell.add_paragraph()
geo_run = geo_score.add_run('GEO健康评分：')
geo_run.font.size = Pt(12)
geo_run.font.bold = True
score_run = geo_score.add_run('68分')
score_run.font.size = Pt(14)
score_run.font.bold = True
score_run.font.color.rgb = RGBColor(230, 126, 34)  # 橙色醒目
level_run = geo_score.add_run(' | ⭐⭐ C级')
level_run.font.size = Pt(12)
geo_score.paragraph_format.space_after = Pt(6)

# 其他评分项
vis_score = cell.add_paragraph()
vis_score.add_run('可见度评分：').font.bold = True
vis_run = vis_score.add_run('72分')
vis_run.font.bold = True
vis_run.font.color.rgb = RGBColor(230, 126, 34)
vis_score.add_run(' | 较好可见性')
vis_score.paragraph_format.space_after = Pt(6)

rec_score = cell.add_paragraph()
rec_score.add_run('推荐度评分：').font.bold = True
rec_run = rec_score.add_run('65分')
rec_run.font.bold = True
rec_run.font.color.rgb = RGBColor(230, 126, 34)
rec_score.add_run(' | 推荐意愿中等')
rec_score.paragraph_format.space_after = Pt(6)

sent_score = cell.add_paragraph()
sent_score.add_run('正面评价占比：').font.bold = True
sent_run = sent_score.add_run('58%')
sent_run.font.bold = True
sent_run.font.color.rgb = RGBColor(230, 126, 34)
sent_score.add_run(' | 整体舆情尚可')

# 设置表格背景色
set_cell_shading(cell, 'F5F5F5')

doc.add_paragraph().paragraph_format.space_after = Pt(6)

# 核心优势
adv_title = doc.add_paragraph()
add_run_with_color(adv_title, '核心优势：', bold=True)
adv_title.paragraph_format.space_after = Pt(4)

advantages = [
    '• 武汉老牌高端商业地标',
    '• 武商集团核心旗舰项目',
    '• 奢侈品品牌资源丰富'
]
for adv in advantages:
    p = doc.add_paragraph(adv)
    p.paragraph_format.left_indent = Inches(0.3)
    p.paragraph_format.space_after = Pt(2)

doc.add_paragraph().paragraph_format.space_after = Pt(6)

# 主要挑战
chall_title = doc.add_paragraph()
add_run_with_color(chall_title, '主要挑战：', bold=True)
chall_title.paragraph_format.space_after = Pt(4)

challenges = [
    '• 数字化内容布局不足',
    '• 与新兴竞品差异化定位模糊',
    '• 场景化内容输出较少'
]
for chall in challenges:
    p = doc.add_paragraph(chall)
    p.paragraph_format.left_indent = Inches(0.3)
    p.paragraph_format.space_after = Pt(2)

# 分隔线
line2 = doc.add_paragraph()
add_horizontal_line(line2)
line2.paragraph_format.space_after = Pt(12)

# ========== 一、AI搜索表现 ==========
sec1_title = doc.add_paragraph()
add_run_with_color(sec1_title, '🔍 一、AI搜索表现', size=16, bold=True)
sec1_title.paragraph_format.space_after = Pt(12)

sub1_title = doc.add_paragraph()
add_run_with_color(sub1_title, '1.1 品牌可见度分析', bold=True)
sub1_title.paragraph_format.space_after = Pt(6)

p1_1 = doc.add_paragraph()
p1_1.add_run('在本次AI搜索分析中，武汉国际广场累计获得')
add_run_with_color(p1_1, '45次', bold=True, color=(230, 126, 34))
p1_1.add_run('AI平台提及。')
p1_1.paragraph_format.space_after = Pt(8)

vis_intro = doc.add_paragraph()
vis_intro.add_run('整体可见度得分：')
add_run_with_color(vis_intro, '72/100', bold=True, size=13, color=(230, 126, 34))
vis_intro.paragraph_format.space_after = Pt(8)

tips_title = doc.add_paragraph()
add_run_with_color(tips_title, '可见度提升关键点：', bold=True)
tips_title.paragraph_format.space_after = Pt(4)

tips = [
    '① 强化在小红书、抖音等内容平台的种草内容',
    '② 建立品牌核心关键词与场景的强关联',
    '③ 增加知乎、大众点评等平台的专业内容覆盖'
]
for tip in tips:
    p = doc.add_paragraph(tip)
    p.paragraph_format.left_indent = Inches(0.3)
    p.paragraph_format.space_after = Pt(2)

doc.add_paragraph().paragraph_format.space_after = Pt(6)

sub1_2 = doc.add_paragraph()
add_run_with_color(sub1_2, '1.2 平台提及详情', bold=True)
sub1_2.paragraph_format.space_after = Pt(6)

platforms = [
    ('豆包：', '11次'),
    ('DeepSeek：', '10次'),
    ('元宝：', '9次'),
    ('Kimi：', '8次'),
    ('通义千问：', '9次'),
    ('智谱：', '12次'),
]
for name, count in platforms:
    p = doc.add_paragraph()
    p.add_run(name)
    count_run = p.add_run(count)
    count_run.font.bold = True
    count_run.font.color.rgb = RGBColor(230, 126, 34)
    p.paragraph_format.left_indent = Inches(0.3)
    p.paragraph_format.space_after = Pt(2)

doc.add_paragraph().paragraph_format.space_after = Pt(6)

sub1_3 = doc.add_paragraph()
add_run_with_color(sub1_3, '1.3 AI推荐理由', bold=True)
sub1_3.paragraph_format.space_after = Pt(6)

p1_3_intro = doc.add_paragraph('武汉国际广场在AI搜索中的典型推荐理由：')
p1_3_intro.paragraph_format.space_after = Pt(6)

# 引用块
quote_para = doc.add_paragraph()
quote_para.paragraph_format.left_indent = Inches(0.4)
quote_para.paragraph_format.right_indent = Inches(0.2)
quote_run = quote_para.add_run('"武汉国际广场是武商集团的旗舰商业项目，位于解放大道核心商圈，是武汉历史最悠久的高端购物中心之一。商场汇聚了众多国际一线奢侈品牌，业态涵盖购物、餐饮、娱乐休闲等，是追求品质消费的热门目的地。"')
quote_run.font.italic = True
quote_run.font.color.rgb = RGBColor(100, 100, 100)

# 分隔线
line3 = doc.add_paragraph()
add_horizontal_line(line3)
line3.paragraph_format.space_after = Pt(12)

# ========== 二、品牌健康度 ==========
sec2_title = doc.add_paragraph()
add_run_with_color(sec2_title, '💚 二、品牌健康度', size=16, bold=True)
sec2_title.paragraph_format.space_after = Pt(12)

sub2_1 = doc.add_paragraph()
add_run_with_color(sub2_1, '2.1 GEO健康评分', bold=True)
sub2_1.paragraph_format.space_after = Pt(6)

geo_para = doc.add_paragraph()
geo_para.add_run('GEO健康评分：')
geo_score_run = geo_para.add_run('68/100')
geo_score_run.font.size = Pt(13)
geo_score_run.font.bold = True
geo_score_run.font.color.rgb = RGBColor(230, 126, 34)
geo_para.add_run(' | ⭐⭐ C级')
geo_para.paragraph_format.space_after = Pt(8)

p2_1_desc = doc.add_paragraph('品牌在AI搜索中表现一般，具有一定基础但优化空间较大。建议加强内容建设，提升品牌在AI平台的竞争力。')

doc.add_paragraph().paragraph_format.space_after = Pt(6)

sub2_2 = doc.add_paragraph()
add_run_with_color(sub2_2, '2.2 三维度评估', bold=True)
sub2_2.paragraph_format.space_after = Pt(6)

# 维度评分
dimensions = [
    ('可见度：', '72/100', 'B级', '🟡 良好'),
    ('推荐度：', '65/100', 'C级', '🟠 一般'),
    ('情感倾向：', '58/100', 'C级', '🟠 一般'),
]
for dim_name, dim_score, dim_level, dim_status in dimensions:
    p = doc.add_paragraph()
    p.add_run(dim_name)
    score_r = p.add_run(dim_score)
    score_r.font.bold = True
    score_r.font.color.rgb = RGBColor(230, 126, 34)
    p.add_run(f' | {dim_level} · {dim_status}')
    p.paragraph_format.left_indent = Inches(0.3)
    p.paragraph_format.space_after = Pt(4)

doc.add_paragraph().paragraph_format.space_after = Pt(4)

eval_title = doc.add_paragraph()
add_run_with_color(eval_title, '评估说明：', bold=True)
eval_title.paragraph_format.space_after = Pt(4)

evals = [
    '• 可见度：品牌在AI搜索结果中被提及的频率和覆盖范围',
    '• 推荐度：AI在回答相关问题时推荐该品牌的概率',
    '• 情感倾向：品牌相关内容的正面评价占比'
]
for ev in evals:
    p = doc.add_paragraph(ev)
    p.paragraph_format.left_indent = Inches(0.3)
    p.paragraph_format.space_after = Pt(2)

doc.add_paragraph().paragraph_format.space_after = Pt(6)

sub2_3 = doc.add_paragraph()
add_run_with_color(sub2_3, '2.3 健康度趋势', bold=True)
sub2_3.paragraph_format.space_after = Pt(6)

trend_para = doc.add_paragraph()
trend_para.add_run('➡️ 稳中有降趋势')
trend_run = trend_para.add_run(' - 面临新兴竞品冲击，需加强数字化营销')
trend_run.font.italic = True
trend_para.paragraph_format.space_after = Pt(4)

# 分隔线
line4 = doc.add_paragraph()
add_horizontal_line(line4)
line4.paragraph_format.space_after = Pt(12)

# ========== 三、竞品对比 ==========
sec3_title = doc.add_paragraph()
add_run_with_color(sec3_title, '🏢 三、竞品对比', size=16, bold=True)
sec3_title.paragraph_format.space_after = Pt(12)

sub3_1 = doc.add_paragraph()
add_run_with_color(sub3_1, '3.1 竞品概述', bold=True)
sub3_1.paragraph_format.space_after = Pt(6)

p3_1 = doc.add_paragraph('本次分析纳入了3个主要竞品进行对比分析：')
p3_1.paragraph_format.space_after = Pt(8)

# 排名
rankings = [
    ('🥇 第1名：武汉恒隆广场', ' — 82分', True),
    ('🥈 第2名：武商MALL', ' — 78分', True),
    ('🥉 第3名：武汉万象城', ' — 75分', True),
    ('④ 武汉国际广场', ' — 68分', False),
]
for name, score, is_highlight in rankings:
    p = doc.add_paragraph()
    name_run = p.add_run(name)
    if is_highlight:
        name_run.font.bold = True
    score_run = p.add_run(score)
    score_run.font.bold = True
    score_run.font.color.rgb = RGBColor(230, 126, 34)
    p.paragraph_format.left_indent = Inches(0.3)
    p.paragraph_format.space_after = Pt(4)

doc.add_paragraph().paragraph_format.space_after = Pt(6)

sub3_2 = doc.add_paragraph()
add_run_with_color(sub3_2, '3.2 竞争优势识别', bold=True)
sub3_2.paragraph_format.space_after = Pt(6)

adv_sec = doc.add_paragraph()
add_run_with_color(adv_sec, '✅ 差异化竞争优势', bold=True)
adv_sec.paragraph_format.space_after = Pt(4)

adv_items = [
    '① 历史底蕴深厚：作为武汉最早的顶级购物中心，具有不可复制的历史积淀',
    '② 武商集团背书：依托武商集团的品牌影响力和资源整合能力',
    '③ 核心区位优势：占据解放大道黄金商圈核心位置'
]
for item in adv_items:
    p = doc.add_paragraph(item)
    p.paragraph_format.left_indent = Inches(0.3)
    p.paragraph_format.space_after = Pt(2)

doc.add_paragraph().paragraph_format.space_after = Pt(6)

warn_sec = doc.add_paragraph()
add_run_with_color(warn_sec, '⚠️ 劣势分析', bold=True)
warn_sec.paragraph_format.space_after = Pt(4)

warn_items = [
    '① 数字化程度偏低：相比新兴竞品，线上内容运营能力不足',
    '② 场景化内容缺失：缺乏针对约会、遛娃等场景的系统性内容输出',
    '③ 新媒体声量较弱：小红书、抖音等平台的内容质量和更新频率待提升'
]
for item in warn_items:
    p = doc.add_paragraph(item)
    p.paragraph_format.left_indent = Inches(0.3)
    p.paragraph_format.space_after = Pt(2)

# 分隔线
line5 = doc.add_paragraph()
add_horizontal_line(line5)
line5.paragraph_format.space_after = Pt(12)

# ========== 四、优化建议 ==========
sec4_title = doc.add_paragraph()
add_run_with_color(sec4_title, '💡 四、优化建议', size=16, bold=True)
sec4_title.paragraph_format.space_after = Pt(12)

sub4_1 = doc.add_paragraph()
add_run_with_color(sub4_1, '4.1 优先级矩阵', bold=True)
sub4_1.paragraph_format.space_after = Pt(6)

priorities = [
    ('🔴 高优先级：', '强化场景化内容', True),
    ('🔴 高优先级：', '优化品牌咨询内容', True),
    ('🟠 中优先级：', '优化对比内容', False),
    ('🟡 低优先级：', '技术SEO优化', False),
]
for prefix, content, is_high in priorities:
    p = doc.add_paragraph()
    pre_run = p.add_run(prefix)
    pre_run.font.bold = is_high
    cont_run = p.add_run(content)
    cont_run.font.bold = is_high
    p.paragraph_format.left_indent = Inches(0.3)
    p.paragraph_format.space_after = Pt(2)

doc.add_paragraph().paragraph_format.space_after = Pt(6)

sub4_2 = doc.add_paragraph()
add_run_with_color(sub4_2, '4.2 具体优化措施', bold=True)
sub4_2.paragraph_format.space_after = Pt(6)

# 措施1
m1_title = doc.add_paragraph()
add_run_with_color(m1_title, '① 强化场景化内容', bold=True)
m1_title.paragraph_format.space_after = Pt(4)

m1_desc = doc.add_paragraph('针对约会、遛娃、购物等场景创作专属内容')
m1_desc.paragraph_format.space_after = Pt(4)

m1_items = [
    '• 制作场景化短视频（抖音/视频号）',
    '• 发布小红书攻略笔记',
    '• 建立场景化内容标签体系'
]
for item in m1_items:
    p = doc.add_paragraph(item)
    p.paragraph_format.left_indent = Inches(0.3)
    p.paragraph_format.space_after = Pt(2)

doc.add_paragraph().paragraph_format.space_after = Pt(6)

# 措施2
m2_title = doc.add_paragraph()
add_run_with_color(m2_title, '② 优化品牌咨询内容', bold=True)
m2_title.paragraph_format.space_after = Pt(4)

m2_desc = doc.add_paragraph('增强品牌故事和差异化卖点输出')
m2_desc.paragraph_format.space_after = Pt(4)

m2_items = [
    '• 梳理品牌核心优势',
    '• 制作品牌介绍长图文',
    '• 强化武商集团背书传播'
]
for item in m2_items:
    p = doc.add_paragraph(item)
    p.paragraph_format.left_indent = Inches(0.3)
    p.paragraph_format.space_after = Pt(2)

doc.add_paragraph().paragraph_format.space_after = Pt(6)

# 措施3
m3_title = doc.add_paragraph()
add_run_with_color(m3_title, '③ 优化对比内容', bold=True)
m3_title.paragraph_format.space_after = Pt(4)

m3_desc = doc.add_paragraph('制作与竞品的差异化对比内容')
m3_desc.paragraph_format.space_after = Pt(4)

m3_items = [
    '• 策划PK内容选题',
    '• 发布对比评测文章',
    '• 强化历史底蕴和品质定位'
]
for item in m3_items:
    p = doc.add_paragraph(item)
    p.paragraph_format.left_indent = Inches(0.3)
    p.paragraph_format.space_after = Pt(2)

doc.add_paragraph().paragraph_format.space_after = Pt(6)

sub4_3 = doc.add_paragraph()
add_run_with_color(sub4_3, '4.3 风险预警', bold=True)
sub4_3.paragraph_format.space_after = Pt(6)

risk_para = doc.add_paragraph()
risk_run = risk_para.add_run('⚠️ 中风险：新兴高端商业体持续冲击传统高端市场，建议加速数字化转型')

# 分隔线
line6 = doc.add_paragraph()
add_horizontal_line(line6)
line6.paragraph_format.space_after = Pt(12)

# ========== 五、武汉高端商业竞争格局 ==========
sec5_title = doc.add_paragraph()
add_run_with_color(sec5_title, '📝 五、武汉高端商业竞争格局', size=16, bold=True)
sec5_title.paragraph_format.space_after = Pt(12)

# 第一梯队
tier1_title = doc.add_paragraph()
add_run_with_color(tier1_title, '第一梯队', bold=True)
tier1_title.paragraph_format.space_after = Pt(4)

tier1_items = [
    '武汉恒隆广场 — 高端标杆，新生代消费者首选',
    '武商MALL — 综合高端，本地底蕴深厚',
    '武汉万象城 — 品质生活，场景体验丰富'
]
for item in tier1_items:
    p = doc.add_paragraph(item)
    p.paragraph_format.left_indent = Inches(0.3)
    p.paragraph_format.space_after = Pt(2)

doc.add_paragraph().paragraph_format.space_after = Pt(6)

# 第二梯队
tier2_title = doc.add_paragraph()
add_run_with_color(tier2_title, '第二梯队', bold=True)
tier2_title.paragraph_format.space_after = Pt(4)

tier2_items = [
    '武汉国际广场 — 历史地标，品牌积淀深厚',
    '武汉K11 — 艺术特色，差异化定位'
]
for item in tier2_items:
    p = doc.add_paragraph(item)
    p.paragraph_format.left_indent = Inches(0.3)
    p.paragraph_format.space_after = Pt(2)

# 分隔线
line7 = doc.add_paragraph()
add_horizontal_line(line7)
line7.paragraph_format.space_after = Pt(12)

# ========== 附录 ==========
appendix_title = doc.add_paragraph()
add_run_with_color(appendix_title, '📎 附录', size=16, bold=True)
appendix_title.paragraph_format.space_after = Pt(12)

# 数据来源
source_title = doc.add_paragraph()
add_run_with_color(source_title, '数据来源', bold=True)
source_title.paragraph_format.space_after = Pt(4)

sources = [
    '• AI搜索平台：豆包、千问、DeepSeek、元宝、Kimi、智谱AI',
    '• 高权重内容平台：小红书、知乎、微信公众号、抖音、大众点评',
    '• 数据采集时间：2026年05月09日'
]
for s in sources:
    p = doc.add_paragraph(s)
    p.paragraph_format.left_indent = Inches(0.3)
    p.paragraph_format.space_after = Pt(2)

doc.add_paragraph().paragraph_format.space_after = Pt(6)

# 评分标准
std_title = doc.add_paragraph()
add_run_with_color(std_title, '评分标准', bold=True)
std_title.paragraph_format.space_after = Pt(4)

standards = [
    '• S级（90-100分）：极优秀，AI搜索首选推荐',
    '• A级（80-89分）：优秀，AI搜索高频提及',
    '• B级（70-79分）：良好，AI搜索稳定可见',
    '• C级（60-69分）：一般，AI搜索偶尔提及',
    '• D级（<60分）：不足，需要重点优化'
]
for s in standards:
    p = doc.add_paragraph(s)
    p.paragraph_format.left_indent = Inches(0.3)
    p.paragraph_format.space_after = Pt(2)

doc.add_paragraph().paragraph_format.space_after = Pt(6)

# 术语说明
term_title = doc.add_paragraph()
add_run_with_color(term_title, '术语说明', bold=True)
term_title.paragraph_format.space_after = Pt(4)

terms = [
    '• GEO：Generative Engine Optimization，生成式引擎优化',
    '• 可见度：品牌在AI搜索结果中被提及的频率',
    '• 推荐度：AI搜索结果中品牌被推荐的概率',
    '• 情感倾向：品牌相关内容的正面/中性/负面比例'
]
for t in terms:
    p = doc.add_paragraph(t)
    p.paragraph_format.left_indent = Inches(0.3)
    p.paragraph_format.space_after = Pt(2)

# 分隔线
line8 = doc.add_paragraph()
add_horizontal_line(line8)
line8.paragraph_format.space_after = Pt(12)

# ========== 底部信息 ==========
contact = doc.add_paragraph()
contact.alignment = WD_ALIGN_PARAGRAPH.CENTER
contact_run = contact.add_run('📞 商务合作：189 8603 4050')
contact_run.font.size = Pt(12)
contact_run.font.bold = True
contact.paragraph_format.space_after = Pt(8)

credit = doc.add_paragraph()
credit.alignment = WD_ALIGN_PARAGRAPH.CENTER
credit_run = credit.add_run('本报告由 好易易GEO 自动生成')
credit_run.font.size = Pt(10)
credit_run.font.color.rgb = RGBColor(128, 128, 128)
credit_run.font.italic = True

# 保存文档
import os
script_dir = os.path.dirname(os.path.abspath(__file__))
output_path = os.path.join(script_dir, '武汉国际广场_GEO报告_公众号版.docx')
doc.save(output_path)
print(f'✅ 文档已成功生成：{output_path}')
